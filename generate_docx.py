import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    kwargs can be top, bottom, left, right.
    val: 'single', 'double', 'dashed', etc.
    color: 'auto' or hex code
    sz: size in 1/8 pt
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name, border_attrs in kwargs.items():
        border = OxmlElement(f'w:{border_name}')
        for attr, val in border_attrs.items():
            border.set(qn(f'w:{attr}'), str(val))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_callout_box(doc, text_lines, title=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    set_cell_border(cell, 
                    left={'val': 'single', 'sz': 24, 'color': '1A365D'},
                    top={'val': 'none'}, right={'val': 'none'}, bottom={'val': 'none'})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if title:
        run_title = p.add_run(title + "\n")
        run_title.bold = True
        run_title.font.name = "Times New Roman"
        run_title.font.size = Pt(10.5)
        run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
    for i, line in enumerate(text_lines):
        if i > 0 or title:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
        
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    # Empty paragraph after table for spacing
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

print("Helper functions compiled successfully.")
